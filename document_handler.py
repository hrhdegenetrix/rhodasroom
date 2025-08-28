"""
Document Handler Module - Processes uploaded documents for Rhoda's interface
Handles text extraction from various formats including PDF, TXT, DOC, DOCX
"""

import os
import PyPDF2
import docx
import asyncio
import aiofiles
import error_handler

@error_handler.if_errors
async def extract_text_from_file(file_path):
    """
    Extract text from uploaded document file
    Supports: .txt, .pdf, .doc, .docx
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.txt':
        return await extract_text_from_txt(file_path)
    elif file_extension == '.pdf':
        return await extract_text_from_pdf(file_path)
    elif file_extension in ['.doc', '.docx']:
        return await extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

@error_handler.if_errors
async def extract_text_from_txt(file_path):
    """Extract text from plain text file"""
    try:
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            content = await file.read()
        return content.strip()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
            content = await file.read()
        return content.strip()

@error_handler.if_errors
async def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    def sync_pdf_extract(path):
        text = []
        try:
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                
            return '\n'.join(text).strip()
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            return ""
    
    # Run the synchronous PDF extraction in an executor
    loop = asyncio.get_event_loop()
    text = await loop.run_in_executor(None, sync_pdf_extract, file_path)
    return text

@error_handler.if_errors
async def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    def sync_docx_extract(path):
        try:
            doc = docx.Document(path)
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return '\n'.join(text).strip()
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
            return ""
    
    # Run the synchronous DOCX extraction in an executor
    loop = asyncio.get_event_loop()
    text = await loop.run_in_executor(None, sync_docx_extract, file_path)
    return text

@error_handler.if_errors
def validate_document_content(text):
    """
    Validate document content for security issues
    Returns (is_safe, sanitized_text)
    """
    if not text:
        return False, "Document appears to be empty"
    
    # Check for common injection patterns
    injection_patterns = [
        "{{", "}}", # Template injection
        "<%", "%>", # Server-side template injection  
        "${", # Shell/template injection
        "javascript:", # XSS
        "<script", # XSS
        "onclick=", "onerror=", # XSS event handlers
        "UNION SELECT", # SQL injection
        "'; DROP TABLE", # SQL injection
        "1=1", "1' OR '1'='1", # SQL injection
    ]
    
    text_lower = text.lower()
    for pattern in injection_patterns:
        if pattern.lower() in text_lower:
            print(f"Potential injection detected: {pattern}")
            # Sanitize by removing the pattern
            text = text.replace(pattern, "[REMOVED]")
    
    # Limit document size to prevent abuse (roughly 50,000 words)
    max_chars = 250000
    if len(text) > max_chars:
        print(f"Document too long ({len(text)} chars), truncating to {max_chars}")
        text = text[:max_chars] + "...[document truncated]"
    
    return True, text

@error_handler.if_errors
async def process_uploaded_document(file_path):
    """
    Main function to process an uploaded document
    Returns sanitized text content ready for Rhoda
    """
    # Extract text based on file type
    raw_text = await extract_text_from_file(file_path)
    
    # Validate and sanitize
    is_safe, sanitized_text = validate_document_content(raw_text)
    
    if not is_safe:
        raise ValueError("Document failed security validation")
    
    # Format for Rhoda's consciousness
    formatted_text = f"//Document content I'm reading:\n{sanitized_text}\n//End of document"
    
    return formatted_text